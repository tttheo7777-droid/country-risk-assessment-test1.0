"""
国别风险评估系统 - 报告生成器模块
功能：组装完整评估报告，支持Markdown/Word/PDF三种格式导出
"""
import io
from datetime import date
from config import DIMENSION_LABELS, INDICATOR_META, VERSION
from report.templates import (
    generate_executive_summary, generate_dimension_analysis,
    generate_risk_pathways, generate_recommendations,
)


def generate_full_report(country_data, result):
    """生成完整的Markdown格式评估报告。

    Args:
        country_data: get_country_data() 返回的完整国家数据
        result:       calculate_total_score() 返回的评分结果

    Returns:
        str: 完整的Markdown格式报告文本
    """
    name = country_data["name_cn"]
    ename = country_data["name_en"]
    today = date.today().strftime("%Y年%m月%d日")

    lines = [
        f"# 国别海外利益安全风险评估报告",
        f"",
        f"## {name}（{ename}）",
        f"",
        f"---",
        f"",
        f"| 项目 | 内容 |",
        f"|------|------|",
        f"| 评估日期 | {today} |",
        f"| 评估国家 | {name}（{ename}） |",
        f"| 所属区域 | {country_data['region']} |",
        f"| 综合风险总分 | **{result['total_score']}/100** |",
        f"| 风险等级 | **{result['risk_label']}** |",
        f"| 系统版本 | 国别风险评估系统 v{VERSION} |",
        f"| 数据说明 | {country_data.get('data_source_note', '基于国际指标模拟值')} |",
        f"",
        f"---",
        f"",
        f"## 一、执行摘要",
        f"",
        generate_executive_summary(country_data, result),
        f"",
        f"---",
        f"",
        f"## 二、分维度风险评估分析",
        f"",
    ]

    for dim_key in DIMENSION_LABELS:
        lines.append(generate_dimension_analysis(dim_key, country_data, result))

    lines.extend([
        f"---",
        f"",
        f"## 三、风险传导路径分析",
        f"",
    ])

    pathways = generate_risk_pathways(country_data, result)
    for i, pathway in enumerate(pathways, 1):
        lines.append(f"{pathway}")
        lines.append("")

    lines.extend([
        f"---",
        f"",
        f"## 四、应对建议",
        f"",
    ])

    recommendations = generate_recommendations(country_data, result)
    lines.append(f"### 短期建议（0-6个月）")
    lines.append("")
    lines.append(recommendations["short_term"])
    lines.append("")
    lines.append(f"### 中期建议（6-24个月）")
    lines.append("")
    lines.append(recommendations["medium_term"])
    lines.append("")
    lines.append(f"### 长期建议（24个月以上）")
    lines.append("")
    lines.append(recommendations["long_term"])
    lines.append("")

    lines.extend([
        f"---",
        f"",
        f"## 附录：评分方法论",
        f"",
        f"本报告采用层次分析法（AHP）构建国别风险评估指标体系，包含5个一级维度与15个二级子指标。"
        f"原始数据经Min-Max标准化映射至0-100风险分，再通过加权求和与硬性扣分规则计算综合得分。",
        f"",
        f"**国际数据源对标**：",
        f"- 世界银行全球治理指标（WGI）",
        f"- 世界银行世界发展指标（WDI）",
        f"- 国际货币基金组织国际金融统计（IMF IFS）",
        f"- 联合国毒品和犯罪问题办公室（UNODC）",
        f"- 全球恐怖主义指数（GTI）",
        f"- 全球和平指数（GPI）",
        f"- 乌普萨拉冲突数据计划（UCDP）",
        f"- OFAC特别指定国民清单（SDN List）",
        f"- OECD外资规制限制性指数",
        f"",
        f"**风险等级阈值**：低风险（0-40）/ 中风险（41-60）/ 高风险（61-80）/ 极高风险（81-100）",
        f"",
        f"*声明：本报告数据基于国际指标模拟值（2022-2024），仅供学术研究与风险预判参考，不构成投资建议。*",
    ])

    return "\n".join(lines)


def export_markdown(country_data, result):
    """将完整报告导出为Markdown文本（可直接复制）。

    Returns:
        str: Markdown格式报告
    """
    return generate_full_report(country_data, result)


def export_word(country_data, result):
    """将完整报告导出为Word文档（.docx）。

    Args:
        country_data: 国家完整数据
        result:       评分结果

    Returns:
        bytes: Word文档的二进制内容，可直接写入文件或供下载
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        return None

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "SimSun"
    font.size = Pt(11)

    name = country_data["name_cn"]
    ename = country_data["name_en"]
    today = date.today().strftime("%Y年%m月%d日")

    # 标题
    title = doc.add_heading(f"国别海外利益安全风险评估报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading(f"{name}（{ename}）", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 报告信息表
    doc.add_heading("报告信息", level=2)
    info_table = doc.add_table(rows=6, cols=2, style="Light Grid Accent 1")
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("评估日期", today),
        ("评估国家", f"{name}（{ename}）"),
        ("所属区域", country_data["region"]),
        ("综合风险总分", f"{result['total_score']}/100"),
        ("风险等级", result["risk_label"]),
        ("数据说明", "基于国际指标模拟值（2022-2024）"),
    ]
    for i, (k, v) in enumerate(info_data):
        info_table.cell(i, 0).text = k
        info_table.cell(i, 1).text = v

    doc.add_paragraph()

    # 一、执行摘要
    doc.add_heading("一、执行摘要", level=1)
    summary = generate_executive_summary(country_data, result)
    for para_text in summary.split("\n\n"):
        para_text = para_text.strip()
        if para_text:
            doc.add_paragraph(para_text)

    # 二、分维度分析
    doc.add_heading("二、分维度风险评估分析", level=1)
    for dim_key, dim_label in DIMENSION_LABELS.items():
        score = result["dimension_scores"].get(dim_key, 0)
        doc.add_heading(f"{dim_label}（得分: {score}/100）", level=2)

        subs = result["sub_scores"].get(dim_key, {})
        meta = INDICATOR_META.get(dim_key, {})
        sub_table = doc.add_table(rows=len(subs) + 1, cols=2, style="Light Grid Accent 1")
        sub_table.cell(0, 0).text = "子指标"
        sub_table.cell(0, 1).text = "得分"
        for i, (si_key, si_val) in enumerate(subs.items(), 1):
            si_name = meta.get(si_key, {}).get("name", si_key)
            sub_table.cell(i, 0).text = si_name
            sub_table.cell(i, 1).text = str(si_val)

        penalties = result.get("penalty_details", {}).get(dim_key, [])
        if penalties:
            p_para = doc.add_paragraph()
            p_para.add_run("扣分项: ").bold = True
            for p in penalties:
                doc.add_paragraph(f"  - {p}", style="List Bullet")

        doc.add_paragraph()

    # 三、风险传导路径
    doc.add_heading("三、风险传导路径分析", level=1)
    pathways = generate_risk_pathways(country_data, result)
    for pathway in pathways:
        doc.add_paragraph(pathway, style="List Bullet")

    # 四、应对建议
    doc.add_heading("四、应对建议", level=1)
    recommendations = generate_recommendations(country_data, result)
    for period, title in [("short_term", "短期建议（0-6个月）"),
                           ("medium_term", "中期建议（6-24个月）"),
                           ("long_term", "长期建议（24个月以上）")]:
        doc.add_heading(title, level=2)
        for line in recommendations[period].split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

    # 附录
    doc.add_heading("附录：评分方法论", level=1)
    doc.add_paragraph(
        "本报告采用层次分析法（AHP）构建国别风险评估指标体系，"
        "包含5个一级维度与15个二级子指标。原始数据经Min-Max标准化映射至0-100风险分，"
        "再通过加权求和与硬性扣分规则计算综合得分。"
    )
    doc.add_paragraph(
        "风险等级阈值：低风险（0-40）/ 中风险（41-60）/ 高风险（61-80）/ 极高风险（81-100）。"
    )
    doc.add_paragraph(
        "声明：本报告数据基于国际指标模拟值（2022-2024），仅供学术研究与风险预判参考，不构成投资建议。"
    )

    # 写入内存buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def export_pdf(country_data, result):
    """将完整报告导出为PDF（通过Markdown转PDF的占位实现）。

    Args:
        country_data: 国家完整数据
        result:       评分结果

    Returns:
        bytes | None: PDF内容，若streamlit-pdf不可用则返回None
    """
    return None  # PDF导出需结合streamlit-pdf或weasyprint等外部库
